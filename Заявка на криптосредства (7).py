#!/usr/bin/env python3
"""
Генератор заявок ViPNet PKI Client
Установка: pip3 install python-docx openpyxl xlrd
Запуск:    python3 pki_zayavka.py
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import json, os, subprocess, sys, copy, zipfile, re, socket
from datetime import datetime
import xlrd
import openpyxl
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CONFIG_FILE = os.path.join(os.path.expanduser("~"), ".pki_zayavka_config.json")

MONTHS_RU = {
    1:"января",2:"февраля",3:"марта",4:"апреля",
    5:"мая",6:"июня",7:"июля",8:"августа",
    9:"сентября",10:"октября",11:"ноября",12:"декабря"
}

SMALL_WORDS = {
    'и','или','а','но','для','с','на','по','в','из','до','при',
    'за','над','под','об','о','к','у','от','то','же','бы'
}

PURPOSE_OPTIONS = ["для подписания ЭП на портале ЕЦП"]
PURPOSE_CUSTOM  = "— ввести свой вариант —"

PC_HEADERS = [
    "этаж", "№ кабинета", "Имя ПК", "доменное имя", "ip-адрес",
    "ФИО пользователя", "срок действия серт.", "СКЗИ",
    "Серийный номер на корпусе", "Инвентарный номер", "№ стикера", ""
]
NUM_PC_COLS = len(PC_HEADERS)  # 12


# ══════════════════════════════════════════════════════
#  Тема Excel → реальные цвета
# ══════════════════════════════════════════════════════

# Стандартная схема цветов Office (из theme1.xml файла актуализации)
# Порядок: dk1, lt1, dk2, lt2, accent1-6, hlink, folHlink
THEME_COLORS = [
    '#000000',  # 0  dk1  (windowText)
    '#FFFFFF',  # 1  lt1  (window)
    '#1F497D',  # 2  dk2
    '#EEECE1',  # 3  lt2
    '#4F81BD',  # 4  accent1
    '#C0504D',  # 5  accent2
    '#9BBB59',  # 6  accent3
    '#8064A2',  # 7  accent4
    '#4BACC6',  # 8  accent5
    '#F79646',  # 9  accent6  (оранжевый, с tint → коричневый)
    '#0000FF',  # 10 hlink
    '#800080',  # 11 folHlink
]


def _load_theme_from_xlsx(path: str) -> list:
    """Пытается извлечь цвета темы прямо из xlsx. Возвращает THEME_COLORS если не получилось."""
    try:
        with zipfile.ZipFile(path) as z:
            theme_files = [f for f in z.namelist() if re.search(r'theme/theme\d*\.xml$', f)]
            if not theme_files:
                return THEME_COLORS
            content = z.read(theme_files[0]).decode('utf-8')
        # Порядок тегов в clrScheme: dk1, lt1, dk2, lt2, accent1..6, hlink, folHlink
        colors = []
        for m in re.finditer(r'<a:(?:dk1|lt1|dk2|lt2|accent\d|hlink|folHlink)>'
                             r'.*?(?:<a:srgbClr val="([0-9A-Fa-f]{6})"'
                             r'|<a:sysClr[^>]+lastClr="([0-9A-Fa-f]{6})")',
                             content):
            rgb = m.group(1) or m.group(2)
            colors.append('#' + rgb.upper())
        return colors if len(colors) >= 10 else THEME_COLORS
    except Exception:
        return THEME_COLORS


def _apply_tint(hex_color: str, tint: float) -> str:
    """Применяет tint к hex-цвету: tint<0 темнее, tint>0 светлее."""
    if not hex_color or tint == 0.0:
        return hex_color
    hex_color = hex_color.lstrip('#')
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    if tint < 0:
        r = int(r * (1 + tint))
        g = int(g * (1 + tint))
        b = int(b * (1 + tint))
    else:
        r = int(r + (255 - r) * tint)
        g = int(g + (255 - g) * tint)
        b = int(b + (255 - b) * tint)
    return f'#{max(0,min(255,r)):02X}{max(0,min(255,g)):02X}{max(0,min(255,b)):02X}'


def _resolve_color(color_obj, theme_palette: list) -> str:
    """Превращает openpyxl Color в '#RRGGBB' или '' если нет цвета."""
    if color_obj is None:
        return ''
    try:
        if color_obj.type == 'rgb':
            argb = color_obj.rgb  # 'FFRRGGBB' или '00000000'
            if argb in ('00000000', ''):
                return ''
            return '#' + argb[2:].upper()
        elif color_obj.type == 'theme':
            idx = color_obj.theme
            tint = color_obj.tint or 0.0
            if idx < len(theme_palette):
                return _apply_tint(theme_palette[idx], tint)
        elif color_obj.type == 'indexed':
            # Стандартная indexed палитра (первые 8 — основные)
            INDEXED = ['#000000','#FFFFFF','#FF0000','#00FF00','#0000FF',
                       '#FFFF00','#FF00FF','#00FFFF']
            idx = color_obj.indexed
            if idx < len(INDEXED):
                return INDEXED[idx]
    except Exception:
        pass
    return ''


def _readable_fg(bg: str, fg: str) -> str:
    """Возвращает читаемый цвет текста. Не меняет цвет если фон тёмный."""
    if not fg:
        return '#000000'
    # Определяем яркость фона (пустой bg = белый/светлый)
    if bg:
        r = int(bg[1:3], 16); g = int(bg[3:5], 16); b = int(bg[5:7], 16)
        luminance = 0.299*r + 0.587*g + 0.114*b
        bg_is_dark = luminance < 140
    else:
        bg_is_dark = False  # нет заливки = светлый фон

    # Белый и жёлтый текст нечитаем на светлом фоне
    if not bg_is_dark and fg.upper() in ('#FFFFFF', '#FFFF00'):
        return '#000000'
    return fg


# ══════════════════════════════════════════════════════
#  Конфиг
# ══════════════════════════════════════════════════════

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_config(cfg):
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)


# ══════════════════════════════════════════════════════
#  Форматирование
# ══════════════════════════════════════════════════════

def make_initials(full_name: str) -> str:
    parts = full_name.strip().split()
    if len(parts) >= 2:
        return parts[0] + ' ' + ''.join(p[0] + '.' for p in parts[1:])
    return full_name


def abbreviate_dept(name: str) -> str:
    if not name:
        return name
    parts = name.split()
    result = []
    i = 0
    while i < len(parts):
        w = parts[i]
        if w == '№' and i + 1 < len(parts):
            result.append(' №' + parts[i + 1]); i += 2; continue
        if w.isdigit():
            result.append(w)
        elif '-' not in w and w.lower() in SMALL_WORDS:
            result.append(w.lower())
        elif w[0].isalpha():
            for sp in w.split('-'):
                if sp and sp[0].isalpha():
                    result.append(sp[0].upper())
        i += 1
    return ''.join(result)


def build_position_doc(position: str, department: str) -> str:
    if not department:
        return position
    pw = position.split(); dw = department.split()
    if pw and dw:
        if pw[-1].lower()[:5] == dw[0].lower()[:5]:
            return position + " " + " ".join(dw[1:])
        else:
            return position + " " + department[0].lower() + department[1:]
    return position + " " + department


def chief_position_prefix(chief_position: str) -> str:
    if not chief_position: return ''
    dept_kw = {'отдела','управления','службы','сектора','группы','центра'}
    result = []
    for w in chief_position.split():
        result.append(w)
        if w.lower() in dept_kw: break
    return " ".join(result)


def _strip_suffix(name: str) -> str:
    """'Мошкова Елена Владимировна (ноут)' → 'Мошкова Елена Владимировна'"""
    return re.sub(r'\s*\([^)]*\)\s*$', '', name).strip()


# ══════════════════════════════════════════════════════
#  Загрузка данных
# ══════════════════════════════════════════════════════

def load_phone_book(path: str) -> dict:
    wb = xlrd.open_workbook(path)
    sh = wb.sheets()[0]
    result = {}
    current_dept = ""
    dept_employees = []

    def flush_dept():
        chief = None
        for emp in dept_employees:
            pos = emp['position'].lower()
            if 'начальник' in pos and 'заместитель' not in pos and 'помощник' not in pos:
                chief = emp; break
        if chief is None and dept_employees:
            chief = dept_employees[0]
        for emp in dept_employees:
            emp['chief_name']     = chief['name']                if chief else ''
            emp['chief_position'] = chief['position']            if chief else ''
            emp['chief_initials'] = make_initials(chief['name']) if chief else ''
            result[emp['name']] = emp

    for i in range(sh.nrows):
        row = sh.row_values(i)
        name = str(row[0]).strip()
        pos  = str(row[2]).strip() if len(row) > 2 else ''
        if not name: continue
        if not pos:
            flush_dept(); dept_employees = []; current_dept = name
        else:
            dept_employees.append({'name': name, 'position': pos, 'department': current_dept})
    flush_dept()
    return result


def load_pc_data(path: str) -> dict:
    """
    Возвращает dict: canonical_name → list[record].
    canonical_name — имя БЕЗ суффикса в скобках.
    Каждый record содержит:
      'values'     : list[str] — 11 значений колонок
      'colors'     : list[(bg,fg)] — точные #RRGGBB цвета (тема уже разрешена)
      'display_name': str — ФИО как написано в файле (с суффиксом типа "(ноут)")
      'serial', 'inventory', 'label'
    """
    theme_palette = _load_theme_from_xlsx(path)
    wb = openpyxl.load_workbook(path)
    sh = wb.active
    result = {}

    for row in sh.iter_rows(min_row=2, values_only=False):
        if len(row) < 6 or not row[5].value:
            continue
        display_name = str(row[5].value).strip()
        if not display_name:
            continue
        canonical = _strip_suffix(display_name)

        values = []
        colors = []
        for i in range(NUM_PC_COLS):
            cell = row[i] if i < len(row) else None
            if cell is None:
                values.append(''); colors.append(('', '')); continue

            val = cell.value
            if isinstance(val, datetime):
                val = val.strftime('%d.%m.%Y')
            values.append(str(val).strip() if val is not None else '')

            # Фон: только если явная заливка (patternType='solid')
            has_fill = (cell.fill and cell.fill.patternType == 'solid')
            bg = _resolve_color(cell.fill.fgColor if has_fill else None, theme_palette)
            # #000000 и #FFFFFF — дефолтные цвета темы, не реальная заливка
            if bg in ('#000000', '#FFFFFF'): bg = ''

            # Текст: берём цвет если он явно задан и не чёрный по умолчанию
            fg = _resolve_color(cell.font.color if cell.font else None, theme_palette)
            if fg == '#000000': fg = ''  # чёрный текст = дефолт, не показываем

            colors.append((bg, fg))

        serial    = values[8]
        inventory = values[9]
        cabinet   = values[1]
        label = f"с/н: {serial}  |  инв: {inventory}"
        if cabinet: label += f"  [{cabinet}]"
        if display_name != canonical:
            label += f"  {display_name[len(canonical):]}"  # показываем суффикс "(ноут)"

        result.setdefault(canonical, []).append({
            'values':       values,
            'colors':       colors,
            'display_name': display_name,
            'serial':       serial,
            'inventory':    inventory,
            'label':        label,
        })

    return result


def get_journal_info(path: str) -> dict:
    doc = Document(path)
    t   = doc.tables[0]
    last_pp = 0; last_reg = ''; last_row_idx = 1; last_executor = ''
    for i, row in enumerate(t.rows):
        vals = [c.text.strip() for c in row.cells]
        if vals[0].isdigit():
            last_pp = int(vals[0]); last_reg = vals[2]; last_row_idx = i
            if len(vals) > 4 and vals[4]: last_executor = vals[4]
    next_reg = last_reg
    if last_reg and '-' in last_reg:
        prefix, num = last_reg.rsplit('-', 1)
        try: next_reg = f"{prefix}-{int(num) + 1}"
        except ValueError: pass
    return {
        'next_pp': last_pp + 1, 'next_reg': next_reg,
        'last_row_idx': last_row_idx, 'last_executor': last_executor,
    }


# ══════════════════════════════════════════════════════
#  Журнал
# ══════════════════════════════════════════════════════

def _set_cell_fmt(cell, text, font_name='Times New Roman', font_size=10):
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run(text)
    run.font.name = font_name; run.font.size = Pt(font_size)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'): rFonts.set(qn(attr), font_name)
    ex = rPr.find(qn('w:rFonts'))
    if ex is not None: rPr.remove(ex)
    rPr.insert(0, rFonts)


def add_journal_entry(path: str, entry: dict, last_row_idx: int):
    doc = Document(path); t = doc.tables[0]
    src_tr = t.rows[last_row_idx]._tr
    new_tr = copy.deepcopy(src_tr)
    for tc in new_tr.findall(qn('w:tc')):
        for p in tc.findall(qn('w:p')): tc.remove(p)
        tc.append(OxmlElement('w:p'))
    src_tr.addnext(new_tr)
    new_row = next((r for r in t.rows if r._tr is new_tr), None)
    if new_row is None: raise RuntimeError("Не удалось вставить строку в журнал")
    values = [str(entry['pp']), entry['date'], entry['reg'],
              entry['description'], entry['executor'], entry.get('note', '')]
    for i, val in enumerate(values[:len(new_row.cells)]):
        _set_cell_fmt(new_row.cells[i], val)
    doc.save(path)


# ══════════════════════════════════════════════════════
#  Генерация Word
# ══════════════════════════════════════════════════════

def _run(para, text, bold=False, underline=False, italic=False,
         size=12, name='Times New Roman', superscript=False):
    run = para.add_run(text)
    run.bold = bold; run.underline = underline; run.italic = italic
    run.font.size = Pt(size); run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    for attr in ('w:ascii','w:hAnsi','w:cs'): rFonts.set(qn(attr), name)
    ex = rPr.find(qn('w:rFonts'))
    if ex is not None: rPr.remove(ex)
    rPr.insert(0, rFonts)
    if superscript:
        va = OxmlElement('w:vertAlign'); va.set(qn('w:val'), 'superscript'); rPr.append(va)
    return run


def _para(container, align=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=0, space_after=0, first_indent_cm=None):
    para = container.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    if first_indent_cm is not None:
        para.paragraph_format.first_line_indent = Cm(first_indent_cm)
    return para


def _set_col_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = tcPr.find(qn('w:tcW'))
    if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(dxa)); tcW.set(qn('w:type'), 'dxa')


def _no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    ex = tblPr.find(qn('w:tblBorders'))
    if ex is not None: tblPr.remove(ex)
    tb = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}'); b.set(qn('w:val'),'none'); b.set(qn('w:sz'),'0'); tb.append(b)
    tblPr.append(tb)


def _cell_border(cell, sides: dict):
    tcPr = cell._tc.get_or_add_tcPr()
    tb = tcPr.find(qn('w:tcBorders'))
    if tb is None: tb = OxmlElement('w:tcBorders'); tcPr.append(tb)
    for side, sz in sides.items():
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single' if sz else 'none')
        el.set(qn('w:sz'), str(sz)); el.set(qn('w:space'),'0'); el.set(qn('w:color'),'000000')
        tb.append(el)


def _no_cell_borders(cell):
    _cell_border(cell, {s: 0 for s in ['top','left','bottom','right']})


def generate_zayavka(data: dict, output_path: str):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(0.89); sec.bottom_margin = Cm(2.88)

    footer = sec.footer
    for p in footer.paragraphs: p.clear()
    fp1 = footer.paragraphs[0]
    fp1.paragraph_format.space_before = Pt(0); fp1.paragraph_format.space_after = Pt(0)
    _run(fp1, f'№ {data["reg_number"]}', italic=True)
    fp2 = footer.add_paragraph()
    fp2.paragraph_format.space_before = Pt(0); fp2.paragraph_format.space_after = Pt(0)
    _run(fp2, data['date_short'], italic=True)

    t0 = doc.add_table(rows=1, cols=3); _no_table_borders(t0)
    _set_col_width(t0.cell(0,0), 3544); _set_col_width(t0.cell(0,1), 1417); _set_col_width(t0.cell(0,2), 3969)
    for ci in range(3): _no_cell_borders(t0.cell(0, ci))
    rc = t0.cell(0, 2)
    for idx, (txt, bold) in enumerate([
        ("РАЗРЕШАЮ", True),
        ("Заместитель управляющего Отделением СФР по", False),
        ("Санкт-Петербургу и Ленинградской", False),
        ("области", False),
        ("_______________        Г.Г. Щемелев", False),
        ("«___» _______ 20 ___ г.", False),
    ]):
        p = rc.paragraphs[0] if idx == 0 else rc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        _run(p, txt, bold=bold)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6); _run(p, "ЗАЯВКА", bold=True)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER); _run(p, "НА ОБУЧЕНИЕ РАБОТЕ И УСТАНОВКУ  КРИПТОСРЕДСТВА")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc, first_indent_cm=1.5)
    _run(p, "Прошу Вас разрешить подготовку к самостоятельной работе с криптографическим средством защиты информации:")

    t1 = doc.add_table(rows=2, cols=1); _no_table_borders(t1)
    _set_col_width(t1.cell(0,0), 9637); _set_col_width(t1.cell(1,0), 9637)
    _no_cell_borders(t1.cell(0,0)); _cell_border(t1.cell(1,0), {'top':4,'bottom':4})
    p = t1.cell(0,0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    _run(p, "СКЗИ  ViPNet CSP в составе ПО ViPNet PKI Client")
    p = t1.cell(1,0).paragraphs[0]
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(173.5)
    p.paragraph_format.first_line_indent = Pt(6.5)
    _run(p, "     (тип криптосредства)", size=8)

    p = _para(doc, first_indent_cm=1.5); _run(p, "работника:")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER); _run(p, data['employee_position_full'] + ",  ")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER); _run(p, data['employee_name'])

    p = _para(doc)
    if data['with_install']:
        _run(p, "с установкой", bold=True, underline=True)
        _run(p, " / ", italic=True); _run(p, "без установки", italic=True)
    else:
        _run(p, "с установкой"); _run(p, " / "); _run(p, "без установки", bold=True, underline=True)
    _run(p, " криптосредства на его рабочее место.")
    p = _para(doc)
    _run(p, "                        ", size=12, superscript=True)
    _run(p, " ",                        size=12, superscript=True)
    _run(p, "(нужное подчеркнуть)",     size=14, superscript=True)

    p = _para(doc, first_indent_cm=1.5); _run(p, "Серийный/ инвентарный номер системного блока:")
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER); _run(p, f"с/н: {data['serial']}  инв. {data['inventory']}")

    p = _para(doc, first_indent_cm=1.5); _run(p, "Необходимость установки средства криптозащиты обусловлена:")
    t2 = doc.add_table(rows=1, cols=1); _no_table_borders(t2)
    _set_col_width(t2.cell(0,0), 9569); _cell_border(t2.cell(0,0), {'bottom':4})
    p = t2.cell(0,0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    _run(p, data.get('purpose', 'для подписания ЭП на портале ЕЦП'))
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER); _run(p, "(наименование решаемой задачи)", size=8)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    t3 = doc.add_table(rows=1, cols=2); _no_table_borders(t3)
    _set_col_width(t3.cell(0,0), 5778); _set_col_width(t3.cell(0,1), 4077)
    for ci in range(2): _no_cell_borders(t3.cell(0, ci))
    lc = t3.cell(0,0); p = lc.paragraphs[0]
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    pos_prefix = data.get('chief_pos_prefix',''); abbrev = data.get('chief_abbrev','')
    if pos_prefix and abbrev:
        _run(p, pos_prefix + "             "); _run(p, abbrev, underline=True)
    elif abbrev:
        _run(p, abbrev, underline=True)
    else:
        _run(p, data.get('chief_position_doc',''))
    p2 = lc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
    _run(p2, "                                         (структурного подразделения Отделения)", size=9)
    rc = t3.cell(0,1); p = rc.paragraphs[0]
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    _run(p, "__________/", size=14); _run(p, data['chief_initials'], underline=True)
    p2 = rc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
    _run(p2, "подпись", size=9)
    _run(p2, "                         ", size=12)
    _run(p2, "фамилия и инициалы", size=9)

    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc); _run(p, f"«{data['day']}» {data['month']} {data['year']} г.", underline=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    p = _para(doc); _run(p, "СОГЛАСОВАНО", bold=True)
    for line in ["Начальником отдела","организационно-технической и ","криптографической защиты информации"]:
        p = _para(doc); _run(p, line)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    p = _para(doc); _run(p, "__________________________    ")
    p = _para(doc)
    _run(p, "(подпись)", size=9)
    _run(p, "             ", size=12)
    _run(p, "(фамилия, инициалы)", size=9)
    p = _para(doc); _run(p, "«___» ______________ 20 ___ г.")

    doc.save(output_path)


# ══════════════════════════════════════════════════════
#  Lock-файл для защиты журнала от одновременной записи
# ══════════════════════════════════════════════════════

LOCK_TTL = 60  # секунд, после которых lock считается устаревшим


def _journal_lock_path(journal_path: str) -> str:
    return journal_path + '.lock'


def _read_lock_info(lock_path: str) -> str:
    """Возвращает строку 'ФИО (ИМЯ_ПК)' из lock-файла."""
    try:
        with open(lock_path, 'r', encoding='utf-8') as f:
            lines = f.read().splitlines()
        hostname = lines[0].strip() if lines else 'неизвестный'
        fio      = lines[2].strip() if len(lines) > 2 else ''
        return f"{fio} ({hostname})" if fio else hostname
    except Exception:
        return 'неизвестный'


def _lock_age_seconds(lock_path: str) -> float:
    try:
        mtime = os.path.getmtime(lock_path)
        return (datetime.now() - datetime.fromtimestamp(mtime)).total_seconds()
    except Exception:
        return LOCK_TTL + 1  # считаем устаревшим если не смогли прочитать


def acquire_journal_lock(journal_path: str, executor_name: str = '') -> tuple:
    """Пытается захватить блокировку журнала.
    Возвращает (True, '') при успехе,
               (False, 'ФИО (ИМЯ_ПК)') если заблокировано другим пользователем,
               (False, None) если не удалось создать lock-файл."""
    lock_path = _journal_lock_path(journal_path)
    if os.path.exists(lock_path):
        age = _lock_age_seconds(lock_path)
        if age < LOCK_TTL:
            return False, _read_lock_info(lock_path)
        # Lock устарел — перезахватываем
    try:
        with open(lock_path, 'w', encoding='utf-8') as f:
            f.write(f"{socket.gethostname()}\n{datetime.now().strftime('%d.%m.%Y %H:%M:%S')}\n{executor_name}")
        return True, ''
    except Exception:
        return False, None  # Не удалось создать lock-файл (нет прав?)


def release_journal_lock(journal_path: str):
    """Удаляет lock-файл журнала."""
    lock_path = _journal_lock_path(journal_path)
    try:
        if os.path.exists(lock_path):
            os.remove(lock_path)
    except Exception:
        pass


# ══════════════════════════════════════════════════════
#  Проверка дубликатов в журнале
# ══════════════════════════════════════════════════════

def check_journal_duplicate(journal_path: str, emp_initials: str, serial: str):
    """Ищет запись где в описании (col 3) одновременно есть инициалы и серийный номер.
    Возвращает строку вида 'Запись №5 от 15.01.2026' или None."""
    try:
        doc = Document(journal_path)
        t = doc.tables[0]
        for row in t.rows:
            vals = [c.text.strip() for c in row.cells]
            if len(vals) > 3:
                desc = vals[3]
                if emp_initials and emp_initials in desc:
                    if serial and serial in desc:
                        pp   = vals[0] if vals[0] else '?'
                        date = vals[1] if len(vals) > 1 else ''
                        return f"Запись №{pp} от {date}"
    except Exception:
        pass
    return None


# ══════════════════════════════════════════════════════
#  GUI — Excel Preview Table
# ══════════════════════════════════════════════════════

class ExcelPreviewTable(tk.Frame):
    FONT      = ('TkDefaultFont', 11)
    FONT_BOLD = ('TkDefaultFont', 11, 'bold')
    HDR_BG    = '#2b5797'
    HDR_FG    = '#ffffff'
    EMPTY_BG  = '#f5f5f5'
    ROW_H     = 30
    COL_W     = [70, 90, 145, 170, 110, 185, 108, 170, 160, 160, 90, 36]

    def __init__(self, parent, **kw):
        super().__init__(parent, **kw)
        self._cells = []  # list of tk.Label (data cells)
        self._build()

    def _build(self):
        self._hbar = ttk.Scrollbar(self, orient='horizontal')
        self._hbar.pack(side='bottom', fill='x')
        self._canvas = tk.Canvas(self, height=self.ROW_H * 2 + 4,
                                 bd=0, highlightthickness=0,
                                 xscrollcommand=self._hbar.set,
                                 xscrollincrement=30)
        self._canvas.pack(side='top', fill='both', expand=True)
        self._hbar.config(command=self._canvas.xview)

        self._inner = tk.Frame(self._canvas, bg='#cccccc')
        self._win = self._canvas.create_window((0,0), window=self._inner, anchor='nw')
        self._inner.bind('<Configure>',
            lambda _: self._canvas.configure(scrollregion=self._canvas.bbox('all')))
        self._canvas.bind('<Configure>',
            lambda _: self._canvas.configure(scrollregion=self._canvas.bbox('all')))
        def _on_hscroll(e):
            d = int(-1*(e.delta/120)) if abs(e.delta) >= 120 else (-1 if e.delta > 0 else 1)
            self._canvas.xview_scroll(d, 'units')
            return 'break'  # не передавать событие дальше в bind_all
        self._canvas.bind('<Shift-MouseWheel>', _on_hscroll)

        # Заголовки
        for ci, h in enumerate(PC_HEADERS):
            w = self.COL_W[ci] if ci < len(self.COL_W) else 90
            fr = tk.Frame(self._inner, width=w, height=self.ROW_H,
                          bg='#cccccc'); fr.grid(row=0, column=ci, padx=(0,1), pady=(0,1))
            fr.grid_propagate(False)
            lbl = tk.Label(fr, text=h, font=self.FONT_BOLD,
                           bg=self.HDR_BG, fg=self.HDR_FG,
                           anchor='center', padx=4, wraplength=w-8, justify='center')
            lbl.place(relwidth=1, relheight=1)

        self._show_empty()

    def _show_empty(self):
        for lbl in self._cells:
            try: lbl.master.destroy()
            except: pass
        self._cells = []
        for ci in range(NUM_PC_COLS):
            w = self.COL_W[ci] if ci < len(self.COL_W) else 90
            fr = tk.Frame(self._inner, width=w, height=self.ROW_H,
                          bg='#cccccc'); fr.grid(row=1, column=ci, padx=(0,1), pady=(0,1))
            fr.grid_propagate(False)
            lbl = tk.Label(fr, text='—', font=self.FONT,
                           bg=self.EMPTY_BG, fg='#999999',
                           anchor='center')
            lbl.place(relwidth=1, relheight=1)
            self._cells.append(lbl)
        self._canvas.configure(height=self.ROW_H * 2 + 6)

    def show(self, pc_records: list):
        # Удаляем старые строки данных
        for lbl in self._cells:
            try: lbl.master.destroy()
            except: pass
        self._cells = []

        n = len(pc_records)
        for ri, rec in enumerate(pc_records):
            row_idx = ri + 1
            alt_bg = '#f9f9f9' if ri % 2 == 0 else '#ffffff'
            for ci in range(NUM_PC_COLS):
                val = rec['values'][ci] if ci < len(rec['values']) else ''
                raw_bg, raw_fg = rec['colors'][ci] if ci < len(rec['colors']) else ('','')
                bg = raw_bg if raw_bg else alt_bg
                fg = _readable_fg(bg, raw_fg)

                w = self.COL_W[ci] if ci < len(self.COL_W) else 90
                fr = tk.Frame(self._inner, width=w, height=self.ROW_H,
                              bg='#cccccc')
                fr.grid(row=row_idx, column=ci, padx=(0,1), pady=(0,1))
                fr.grid_propagate(False)
                lbl = tk.Label(fr, text=val, font=self.FONT,
                               bg=bg, fg=fg,
                               anchor='w', padx=4,
                               wraplength=w-8, justify='left')
                lbl.place(relwidth=1, relheight=1)
                self._cells.append(lbl)

        visible_rows = min(n, 5)
        self._canvas.configure(height=self.ROW_H * (visible_rows + 1) + 6)

    def clear(self):
        self._show_empty()


# ══════════════════════════════════════════════════════
#  GUI — App
# ══════════════════════════════════════════════════════

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Генератор заявок ViPNet PKI Client")
        self.geometry("800x1000")
        self.resizable(True, True)
        self.cfg               = load_config()
        self.phone_data        = {}
        self.pc_data           = {}
        self.journal_info      = {}
        self.all_names         = []
        self._pcs              = []
        self._chief_pos_prefix = ''
        self._chief_abbrev     = ''
        self._build_ui()
        self._load_files()

    def _build_ui(self):
        # Главный canvas + scrollbar для вертикального скролла всего окна
        main_canvas  = tk.Canvas(self, borderwidth=0, highlightthickness=0, yscrollincrement=20)
        vbar = ttk.Scrollbar(self, orient='vertical', command=main_canvas.yview)
        main_canvas.configure(yscrollcommand=vbar.set)
        vbar.pack(side='right', fill='y')
        main_canvas.pack(side='left', fill='both', expand=True)

        main = tk.Frame(main_canvas)
        self._main_win = main_canvas.create_window((0,0), window=main, anchor='nw')
        main.bind('<Configure>',
            lambda _: main_canvas.configure(scrollregion=main_canvas.bbox('all')))
        main_canvas.bind('<Configure>',
            lambda e: main_canvas.itemconfig(self._main_win, width=e.width))

        # Привязываем скролл колёсика ко ВСЕМУ приложению
        def _scroll_delta(event):
            if abs(event.delta) >= 120:
                return int(-1 * (event.delta / 120))
            return -1 if event.delta > 0 else 1

        def _on_mousewheel(event):
            if not event.delta:
                return
            main_canvas.yview_scroll(_scroll_delta(event), 'units')

        def _on_shift_mousewheel(event):
            if not event.delta:
                return
            self._preview._canvas.xview_scroll(_scroll_delta(event), 'units')

        self.bind_all('<MouseWheel>', _on_mousewheel)
        self.bind_all('<Shift-MouseWheel>', _on_shift_mousewheel)
        # macOS trackpad двумя пальцами (Linux/X11)
        self.bind_all('<Button-4>',  lambda e: main_canvas.yview_scroll(-1, 'units'))
        self.bind_all('<Button-5>',  lambda e: main_canvas.yview_scroll( 1, 'units'))

        PAD = dict(padx=8, pady=4)

        # ── Файлы ────────────────────────────────────
        ff = ttk.LabelFrame(main, text="  Файлы данных  ", padding=8)
        ff.pack(fill='x', **PAD)
        file_specs = [
            ("phone_book","Телефонный справочник (.xls):",[("XLS","*.xls"),("Все","*.*")]),
            ("pc_file",   "Актуализация ПК (.xlsx):",     [("XLSX","*.xlsx"),("Все","*.*")]),
            ("journal",   "Журнал регистрации (.docx):",  [("DOCX","*.docx"),("Все","*.*")]),
        ]
        self._file_vars = {}
        for ri, (key, label, ftypes) in enumerate(file_specs):
            ttk.Label(ff, text=label, width=30, anchor='w').grid(row=ri, column=0, sticky='w', pady=2)
            var = tk.StringVar(value=self.cfg.get(key,''))
            self._file_vars[key] = var
            ttk.Entry(ff, textvariable=var, width=42).grid(row=ri, column=1, padx=4)
            ttk.Button(ff, text="…", width=3,
                       command=lambda k=key, ft=ftypes: self._pick_file(k,ft)
                       ).grid(row=ri, column=2)
        ttk.Button(ff, text="↺  Загрузить / обновить файлы",
                   command=self._load_files).grid(row=3, column=0, columnspan=3, pady=(6,0))

        # ── Поиск + статус ────────────────────────────
        fs = ttk.LabelFrame(main, text="  Поиск сотрудника  ", padding=8)
        fs.pack(fill='x', **PAD)
        fs.columnconfigure(1, weight=1)

        # Строка поиска
        ttk.Label(fs, text="ФИО:").grid(row=0, column=0, sticky='w')
        self.search_var = tk.StringVar()
        self.search_var.trace('w', self._on_search_change)
        ttk.Entry(fs, textvariable=self.search_var, width=40).grid(row=0, column=1, padx=4, sticky='ew')
        ttk.Button(fs, text="Найти", command=self._do_search).grid(row=0, column=2)

        # Статус сразу под строкой поиска (в области поиска, справа от списка)
        self.status_var  = tk.StringVar(value="")
        self._status_lbl = tk.Label(fs, textvariable=self.status_var,
                                    foreground='#555', wraplength=550,
                                    anchor='w', justify='left',
                                    font=('TkDefaultFont', 11, 'bold'))
        self._status_lbl.grid(row=0, column=3, padx=(12,4), sticky='ew')
        fs.columnconfigure(3, weight=1)

        # Список результатов
        self.lb = tk.Listbox(fs, height=4, width=50, font=('TkDefaultFont', 10))
        self.lb.grid(row=1, column=0, columnspan=2, pady=(4,0), sticky='ew')
        self.lb.bind('<<ListboxSelect>>', self._on_select)
        sb = ttk.Scrollbar(fs, orient='vertical', command=self.lb.yview)
        sb.grid(row=1, column=2, sticky='ns', pady=(4,0))
        self.lb.configure(yscrollcommand=sb.set)

        # ── Excel-превью ─────────────────────────────
        fp = ttk.LabelFrame(main, text="  Данные из таблицы актуализации  ", padding=6)
        fp.pack(fill='x', **PAD)
        self._preview = ExcelPreviewTable(fp)
        self._preview.pack(fill='x', expand=True)
        self._cert_var = tk.StringVar()
        self._cert_lbl = tk.Label(fp, textvariable=self._cert_var,
                                  anchor='w', font=('TkDefaultFont', 10, 'bold'))
        self._cert_lbl.pack(fill='x', padx=4, pady=(4, 0))

        # ── Данные сотрудника ─────────────────────────
        fe = ttk.LabelFrame(main, text="  Данные сотрудника (для документа)  ", padding=8)
        fe.pack(fill='x', **PAD); fe.columnconfigure(1, weight=1)
        self._emp = {}
        for i, (key, label) in enumerate([
            ("emp_name",          "ФИО сотрудника:"),
            ("emp_position_doc",  "Должность и отдел (полная строка):"),
            ("chief_name",        "Начальник (ФИО):"),
            ("chief_initials",    "Инициалы начальника (для подписи):"),
            ("chief_position_doc","Аббревиатура отдела (для подписи):"),
        ]):
            ttk.Label(fe, text=label, anchor='w', width=38).grid(row=i, column=0, sticky='w', pady=2)
            var = tk.StringVar(); self._emp[key] = var
            ttk.Entry(fe, textvariable=var, width=46).grid(row=i, column=1, padx=4, sticky='ew')

        # ── Компьютер ─────────────────────────────────
        fpc = ttk.LabelFrame(main, text="  Компьютер  ", padding=8)
        fpc.pack(fill='x', **PAD); fpc.columnconfigure(1, weight=1)
        ttk.Label(fpc, text="Поиск по номеру:", anchor='w', width=22).grid(row=0, column=0, sticky='w', pady=2)
        self._num_search_var = tk.StringVar()
        ttk.Entry(fpc, textvariable=self._num_search_var, width=40).grid(row=0, column=1, padx=4, sticky='ew')
        ttk.Button(fpc, text="Найти", command=self._search_by_number).grid(row=0, column=2)
        ttk.Label(fpc, text="Выбор ПК:", anchor='w', width=22).grid(row=1, column=0, sticky='w', pady=2)
        self.pc_combo = ttk.Combobox(fpc, width=50, state='readonly')
        self.pc_combo.grid(row=1, column=1, padx=4, sticky='ew')
        self.pc_combo.bind('<<ComboboxSelected>>', self._on_pc_select)
        self._serial_var = tk.StringVar(); self._inv_var = tk.StringVar()
        for ri, (txt, var) in enumerate([("Серийный номер:", self._serial_var),
                                          ("Инвентарный номер:", self._inv_var)], 2):
            ttk.Label(fpc, text=txt, anchor='w', width=22).grid(row=ri, column=0, sticky='w', pady=2)
            ttk.Entry(fpc, textvariable=var, width=40).grid(row=ri, column=1, padx=4, sticky='ew')

        # ── Цель установки ────────────────────────────
        fpur = ttk.LabelFrame(main, text="  Необходимость установки криптозащиты обусловлена  ", padding=8)
        fpur.pack(fill='x', **PAD); fpur.columnconfigure(1, weight=1)
        ttk.Label(fpur, text="Выбор:", anchor='w', width=10).grid(row=0, column=0, sticky='w', pady=2)
        self._purpose_combo_var = tk.StringVar()
        self._purpose_combo = ttk.Combobox(fpur, textvariable=self._purpose_combo_var,
            values=PURPOSE_OPTIONS + [PURPOSE_CUSTOM], state='readonly', width=52)
        self._purpose_combo.grid(row=0, column=1, padx=4, sticky='ew')
        self._purpose_combo.current(0)
        self._purpose_combo.bind('<<ComboboxSelected>>', self._on_purpose_select)
        self._purpose_custom_frame = ttk.Frame(fpur)
        self._purpose_custom_frame.grid(row=1, column=0, columnspan=2, sticky='ew', pady=(4,0))
        self._purpose_custom_frame.columnconfigure(1, weight=1)
        ttk.Label(self._purpose_custom_frame, text="Свой вариант:", anchor='w', width=14
                  ).grid(row=0, column=0, sticky='w')
        self._purpose_custom_var = tk.StringVar()
        ttk.Entry(self._purpose_custom_frame, textvariable=self._purpose_custom_var,
                  width=46).grid(row=0, column=1, padx=4, sticky='ew')
        self._purpose_custom_frame.grid_remove()

        # ── Установка ─────────────────────────────────
        fi = ttk.LabelFrame(main, text="  Установка  ", padding=8)
        fi.pack(fill='x', **PAD)
        self.install_var = tk.BooleanVar(value=True)
        ttk.Radiobutton(fi, text="✓  С установкой",  variable=self.install_var, value=True ).pack(side='left', padx=16)
        ttk.Radiobutton(fi, text="✗  Без установки", variable=self.install_var, value=False).pack(side='left', padx=16)

        # ── Реквизиты ─────────────────────────────────
        fd = ttk.LabelFrame(main, text="  Реквизиты документа  ", padding=8)
        fd.pack(fill='x', **PAD); fd.columnconfigure(1, weight=1)
        self._doc = {}
        now = datetime.now()
        for i, (key, label, default) in enumerate([
            ("reg_number","Рег. номер:",""),
            ("date_short","Дата (дд.мм.гггг):", now.strftime("%d.%m.%Y")),
            ("executor",  "Исполнитель (журнал):",""),
            ("output_dir","Папка сохранения:", os.path.expanduser("~/Desktop")),
        ]):
            ttk.Label(fd, text=label, anchor='w', width=26).grid(row=i, column=0, sticky='w', pady=2)
            var = tk.StringVar(value=default); self._doc[key] = var
            ttk.Entry(fd, textvariable=var, width=44).grid(row=i, column=1, padx=4, sticky='ew')
            if key == 'output_dir':
                ttk.Button(fd, text="…", width=3,
                           command=self._pick_outdir).grid(row=i, column=2)

        ttk.Button(main, text="📄   Создать заявку и добавить в журнал",
                   command=self._generate).pack(pady=10, ipadx=16, ipady=6)

        # Нижний статус (только для общей инфо о загрузке файлов)
        self._load_status_var = tk.StringVar(value="Укажите файлы данных и нажмите «Загрузить»")
        self._load_status_lbl = ttk.Label(main, textvariable=self._load_status_var,
                                          foreground='#555', wraplength=700)
        self._load_status_lbl.pack(pady=(0,12))

    # ── Helpers ───────────────────────────────────────

    def _set_status(self, text, color='gray'):
        clr = {'green':'#1a7a1a','red':'#cc0000','orange':'#b06000','gray':'#555'}
        self.status_var.set(text)
        self._status_lbl.configure(foreground=clr.get(color,'#555'))

    def _set_load_status(self, text, color='gray'):
        clr = {'green':'#2a7a2a','red':'#cc0000','gray':'#555'}
        self._load_status_var.set(text)
        self._load_status_lbl.configure(foreground=clr.get(color, '#555'))

    def _pick_file(self, key, filetypes):
        path = filedialog.askopenfilename(filetypes=filetypes)
        if path:
            self._file_vars[key].set(path); self.cfg[key] = path; save_config(self.cfg)

    def _pick_outdir(self):
        d = filedialog.askdirectory()
        if d: self._doc['output_dir'].set(d)

    def _open_file(self, path):
        try:
            if sys.platform == 'darwin': subprocess.run(['open', path])
            elif sys.platform == 'win32': os.startfile(path)
            else: subprocess.run(['xdg-open', path])
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть: {e}")

    def _on_purpose_select(self, _=None):
        if self._purpose_combo_var.get() == PURPOSE_CUSTOM:
            self._purpose_custom_frame.grid()
        else:
            self._purpose_custom_frame.grid_remove()

    def _get_purpose(self) -> str:
        sel = self._purpose_combo_var.get()
        if sel == PURPOSE_CUSTOM: return self._purpose_custom_var.get().strip()
        return sel

    # ── Load ──────────────────────────────────────────

    def _load_files(self):
        paths  = {k: self._file_vars[k].get().strip() for k in self._file_vars}
        errors = []; loaded = 0

        if paths['phone_book'] and os.path.exists(paths['phone_book']):
            try:
                self.phone_data = load_phone_book(paths['phone_book'])
                self.all_names  = sorted(self.phone_data.keys()); loaded += 1
            except Exception as e: errors.append(f"Справочник: {e}")

        if paths['pc_file'] and os.path.exists(paths['pc_file']):
            try:
                self.pc_data = load_pc_data(paths['pc_file']); loaded += 1
            except Exception as e: errors.append(f"Актуализация ПК: {e}")

        if paths['journal'] and os.path.exists(paths['journal']):
            try:
                self.journal_info = get_journal_info(paths['journal'])
                self._doc['reg_number'].set(self.journal_info['next_reg'])
                if self.journal_info['last_executor']:
                    self._doc['executor'].set(self.journal_info['last_executor'])
                loaded += 1
            except Exception as e: errors.append(f"Журнал: {e}")

        for k in paths: self.cfg[k] = paths[k]
        save_config(self.cfg)

        ji = self.journal_info
        if errors:
            self._set_load_status("⚠ " + " | ".join(errors), 'red')
        else:
            self._set_load_status(
                f"✓ Загружено {loaded}/3 файлов. "
                f"Сотрудников: {len(self.phone_data)}. "
                f"Следующий рег. номер: {ji.get('next_reg','—')}", 'green')

    # ── Search ────────────────────────────────────────

    def _on_search_change(self, *_):
        q = self.search_var.get().strip().lower()
        self.lb.delete(0, tk.END)
        if len(q) < 2: return
        for name in self.all_names:
            if q in name.lower(): self.lb.insert(tk.END, name)

    def _do_search(self): self._on_search_change()

    def _on_select(self, _):
        sel = self.lb.curselection()
        if not sel: return
        name = self.lb.get(sel[0])
        self.search_var.set(name)
        self._fill_employee(name)

    def _fill_employee(self, name: str):
        emp        = self.phone_data.get(name, {})
        position   = emp.get('position', '')
        department = emp.get('department', '')
        abbrev     = abbreviate_dept(department)

        self._chief_pos_prefix = chief_position_prefix(emp.get('chief_position',''))
        self._chief_abbrev     = abbrev

        self._emp['emp_name'].set(name)
        self._emp['emp_position_doc'].set(build_position_doc(position, department))
        self._emp['chief_name'].set(emp.get('chief_name',''))
        self._emp['chief_initials'].set(emp.get('chief_initials',''))
        self._emp['chief_position_doc'].set(abbrev)

        # Ищем ПК по каноническому имени (без суффиксов "(ноут)" и пр.)
        canonical = _strip_suffix(name)
        pcs = self.pc_data.get(canonical, [])
        # Если не нашли по каноническому — пробуем точное
        if not pcs:
            pcs = self.pc_data.get(name, [])
        self._pcs = pcs

        if not pcs:
            self.pc_combo['values'] = ['Не найден в таблице актуализации']
            self.pc_combo.current(0)
            self._serial_var.set(''); self._inv_var.set('')
            self._preview.clear()
            self._cert_var.set('')
            self._set_status("⚠ ПК не найден в таблице актуализации", 'orange')
        else:
            labels = [pc['label'] for pc in pcs]
            self.pc_combo['values'] = labels
            self.pc_combo.current(0)
            self._serial_var.set(pcs[0]['serial'])
            self._inv_var.set(pcs[0]['inventory'])
            self._preview.show(pcs)
            self._update_cert_status(pcs)
            if len(pcs) == 1:
                self._set_status(f"✓ Найден: {name}", 'green')
            else:
                self._set_status(f"⚠ {len(pcs)} устройства — выберите в «Выбор ПК»", 'orange')

    def _update_cert_status(self, pcs: list):
        """Обновляет лейбл статуса сертификата под таблицей превью."""
        if not pcs:
            self._cert_var.set(''); return
        idx = self.pc_combo.current()
        rec = pcs[idx] if 0 <= idx < len(pcs) else pcs[0]
        date_str = rec['values'][6].strip() if len(rec['values']) > 6 else ''
        if not date_str or date_str.lower() in ('нет', '-', 'none', ''):
            self._cert_var.set('— Сертификат не установлен')
            self._cert_lbl.configure(foreground='#888888'); return
        try:
            cert_date = datetime.strptime(date_str, '%d.%m.%Y').date()
        except ValueError:
            self._cert_var.set('— Сертификат не установлен')
            self._cert_lbl.configure(foreground='#888888'); return
        days_left = (cert_date - datetime.now().date()).days
        if days_left < 0:
            self._cert_var.set(f'✗ Сертификат истёк {-days_left} дн. назад ({date_str})')
            self._cert_lbl.configure(foreground='#cc0000')
        elif days_left < 60:
            self._cert_var.set(f'⚠ Сертификат истекает через {days_left} дн. ({date_str})')
            self._cert_lbl.configure(foreground='#b06000')
        else:
            self._cert_var.set(f'✓ Сертификат действителен до {date_str}')
            self._cert_lbl.configure(foreground='#1a7a1a')

    def _search_by_number(self):
        """Поиск по серийному или инвентарному номеру."""
        query = self._num_search_var.get().strip().lower()
        if not query: return
        if not self.pc_data:
            self._set_status("Таблица актуализации не загружена", 'orange'); return
        found = []
        for canonical, recs in self.pc_data.items():
            for rec in recs:
                if query in rec['serial'].lower() or query in rec['inventory'].lower():
                    if canonical not in found:
                        found.append(canonical)
                    break
        if not found:
            self._set_status("Не найдено", 'orange')
        elif len(found) == 1:
            name = found[0]
            self.search_var.set(name)
            self._fill_employee(name)
        else:
            self.lb.delete(0, tk.END)
            for name in found:
                self.lb.insert(tk.END, name)
            self._set_status(f"Найдено {len(found)} совпадений — выберите из списка", 'orange')

    def _on_pc_select(self, _):
        idx = self.pc_combo.current()
        if 0 <= idx < len(self._pcs):
            self._serial_var.set(self._pcs[idx]['serial'])
            self._inv_var.set(self._pcs[idx]['inventory'])
            self._update_cert_status(self._pcs)

    # ── Generate ──────────────────────────────────────

    def _make_filename(self, emp_initials: str, serial: str) -> str:
        return (f"Заявка на обучение СКЗИ  ViPNet CSP в составе "
                f"ПО ViPNet PKI Client {emp_initials} ({serial}).docx")

    def _ask_duplicate(self, path: str, filename: str) -> str:
        win = tk.Toplevel(self); win.title("Файл уже существует")
        win.resizable(False, False); win.grab_set()
        ttk.Label(win, text=f"Заявка уже существует:\n{filename}",
                  wraplength=400, justify='left', padding=12).pack()
        result = tk.StringVar(value='cancel')
        bf = ttk.Frame(win, padding=8); bf.pack()
        def choose(v): result.set(v); win.destroy()
        ttk.Button(bf, text="📂  Открыть существующую",
                   command=lambda: choose('open')  ).grid(row=0,column=0,padx=6,pady=4,sticky='ew')
        ttk.Button(bf, text="🔄  Заменить",
                   command=lambda: choose('replace')).grid(row=0,column=1,padx=6,pady=4,sticky='ew')
        ttk.Button(bf, text="➕  Создать ещё одну",
                   command=lambda: choose('new')   ).grid(row=1,column=0,padx=6,pady=4,sticky='ew')
        ttk.Button(bf, text="✖  Отмена",
                   command=lambda: choose('cancel') ).grid(row=1,column=1,padx=6,pady=4,sticky='ew')
        win.wait_window(); return result.get()

    def _generate(self):
        emp_name = self._emp['emp_name'].get().strip()
        if not emp_name: messagebox.showerror("Ошибка", "Сотрудник не выбран!"); return
        reg = self._doc['reg_number'].get().strip()
        if not reg: messagebox.showerror("Ошибка", "Укажите регистрационный номер!"); return
        date_str = self._doc['date_short'].get().strip()
        try: dt = datetime.strptime(date_str, "%d.%m.%Y")
        except ValueError: messagebox.showerror("Ошибка", "Формат даты: дд.мм.гггг"); return
        purpose = self._get_purpose()
        if not purpose: messagebox.showerror("Ошибка", "Укажите цель установки!"); return

        serial    = self._serial_var.get().strip()
        inventory = self._inv_var.get().strip()
        with_install = self.install_var.get()
        if not serial and not inventory:
            if not messagebox.askyesno("Предупреждение",
                    "Серийный и инвентарный номера не заполнены.\nПродолжить?"): return

        emp_initials = make_initials(emp_name)

        # Проверка дубликата в журнале
        journal_path_check = self._file_vars['journal'].get().strip()
        if journal_path_check and os.path.exists(journal_path_check) and serial:
            dup = check_journal_duplicate(journal_path_check, emp_initials, serial)
            if dup:
                if not messagebox.askyesno("Дубликат в журнале",
                        f"В журнале уже есть заявка на этого сотрудника с этим ПК:\n{dup}\n\n"
                        f"Продолжить и добавить ещё одну запись?"):
                    return

        abbrev_in_ui = self._emp['chief_position_doc'].get().strip()
        data = {
            'reg_number':             reg,
            'date_short':             date_str,
            'employee_name':          emp_name,
            'employee_position_full': self._emp['emp_position_doc'].get().strip(),
            'chief_name':             self._emp['chief_name'].get().strip(),
            'chief_initials':         self._emp['chief_initials'].get().strip(),
            'chief_pos_prefix':       self._chief_pos_prefix,
            'chief_abbrev':           abbrev_in_ui,
            'chief_position_doc':     abbrev_in_ui,
            'serial': serial, 'inventory': inventory,
            'with_install': with_install, 'purpose': purpose,
            'day': str(dt.day), 'month': MONTHS_RU[dt.month], 'year': str(dt.year),
        }

        out_dir = self._doc['output_dir'].get().strip() or os.path.expanduser("~/Desktop")
        os.makedirs(out_dir, exist_ok=True)
        base_name   = self._make_filename(emp_initials, serial)
        output_path = os.path.join(out_dir, base_name)

        if os.path.exists(output_path):
            choice = self._ask_duplicate(output_path, base_name)
            if choice == 'open': self._open_file(output_path); return
            elif choice == 'cancel': return
            elif choice == 'new':
                ts = datetime.now().strftime("%H-%M-%S")
                output_path = os.path.join(
                    out_dir, base_name.replace('.docx', f'_{ts}.docx'))

        try:
            generate_zayavka(data, output_path)
        except Exception as e:
            messagebox.showerror("Ошибка при создании документа", str(e)); return

        journal_path = self._file_vars['journal'].get().strip()
        journal_msg  = ""
        if journal_path and os.path.exists(journal_path) and self.journal_info:
            executor_name = self._doc['executor'].get().strip()
            acquired, blocker = acquire_journal_lock(journal_path, executor_name)
            if not acquired:
                if blocker is None:
                    messagebox.showwarning("Ошибка записи в журнал",
                        "Не удалось заблокировать журнал для записи.\n"
                        "Папка с журналом, возможно, открыта только для чтения.\n"
                        "Обратитесь к администратору.")
                    journal_msg = "⚠ Нет прав на запись рядом с журналом — запись не добавлена."
                else:
                    messagebox.showwarning("Журнал заблокирован",
                        f"Журнал сейчас редактирует другой пользователь:\n{blocker}\n\n"
                        f"Подождите и попробуйте снова.")
                    journal_msg = f"⚠ Журнал заблокирован: {blocker} — запись не добавлена."
            else:
                try:
                    entry = {
                        'pp':          self.journal_info['next_pp'],
                        'date':        date_str, 'reg': reg,
                        'description': f"Заявка на обучение PKI Client {emp_initials} ({serial})",
                        'executor':    self._doc['executor'].get().strip(),
                        'note':        "Акт установки" if with_install else "",
                    }
                    add_journal_entry(journal_path, entry, self.journal_info['last_row_idx'])
                    self.journal_info = get_journal_info(journal_path)
                    self._doc['reg_number'].set(self.journal_info['next_reg'])
                    journal_msg = f"✓ Запись №{entry['pp']} добавлена в журнал."
                except Exception as e:
                    journal_msg = f"⚠ Журнал не обновлён: {e}"
                finally:
                    release_journal_lock(journal_path)
        else:
            journal_msg = "Журнал не указан — запись не добавлена."

        self._set_status(f"✓ Создан: {os.path.basename(output_path)}", 'green')
        if messagebox.askyesno("Готово!",
                               f"Заявка создана:\n{output_path}\n\n{journal_msg}\n\nОткрыть файл?"):
            self._open_file(output_path)


if __name__ == '__main__':
    app = App()
    app.mainloop()
